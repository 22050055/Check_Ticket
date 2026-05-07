# -*- coding: utf-8 -*-
"""Tạo file Word tóm tắt báo cáo đồ án Check Ticket"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Cài đặt lề trang ───────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# ─── Helper functions ────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=13, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(0,32,96)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    sz = {1: 16, 2: 14, 3: 13}.get(level, 13)
    set_font(run, size=sz, bold=True, color=color)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_para(doc, text, indent=0, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=13, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    return p

def add_bullet(doc, text, indent=1.0):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=13)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=12, bold=True, color=(255,255,255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Màu nền header xanh navy
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "002060")
        tcPr.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx+1]
        bg = "DEEAF1" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), bg)
            tcPr.append(shd)
    return table

# ════════════════════════════════════════════════════════════════════
#  TRANG BÌA
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TRƯỜNG ĐẠI HỌC BÌNH DƯƠNG\nKHOA CÔNG NGHỆ THÔNG TIN")
set_font(run, size=14, bold=True, color=(0,32,96))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─────────────────────────────────")
set_font(run, size=13, color=(0,32,96))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BÁO CÁO TÓM TẮT ĐỒ ÁN NGÀNH")
set_font(run, size=18, bold=True, color=(0,32,96))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HỆ THỐNG KIỂM SOÁT RA/VÀO KHU DU LỊCH\nĐA KÊNH TÍCH HỢP NHẬN DIỆN KHUÔN MẶT\nVÀ DASHBOARD PHÂN TÍCH VẬN HÀNH")
set_font(run, size=16, bold=True, color=(192,0,0))

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Sinh viên thực hiện:  NGÔ MẠNH KHANG  —  MSSV: 22050055")
set_font(run, size=13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Giáo viên hướng dẫn:  DƯƠNG ANH TUẤN")
set_font(run, size=13, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Năm thực hiện:  2026")
set_font(run, size=13)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  LỜI CẢM ƠN
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LỜI CẢM ƠN")
set_font(run, size=16, bold=True, color=(0,32,96))

doc.add_paragraph()

add_para(doc, (
    "Lời đầu tiên, em xin gửi lời cảm ơn chân thành nhất đến Ban Giám hiệu trường Đại học Bình Dương "
    "cùng quý Thầy Cô trong Khoa Công nghệ Thông tin đã tận tình giảng dạy, truyền đạt những kiến thức "
    "quý báu và tạo mọi điều kiện thuận lợi nhất cho em trong suốt quá trình học tập và rèn luyện tại trường."
))

add_para(doc, (
    "Đặc biệt, em xin bày tỏ lòng biết ơn sâu sắc đến Thầy Dương Anh Tuấn — người đã trực tiếp hướng dẫn, "
    "tận tình chỉ bảo và dành nhiều thời gian quý báu để giúp đỡ em từ những bước đầu tiên cho đến khi "
    "hoàn thành đồ án này. Những lời khuyên, sự định hướng và khích lệ của Thầy là động lực vô cùng lớn "
    "giúp em vượt qua những khó khăn, thách thức trong quá trình nghiên cứu và triển khai hệ thống."
))

add_para(doc, (
    "Em cũng xin gửi lời cảm ơn đến gia đình và bạn bè đã luôn ở bên động viên, hỗ trợ và đồng hành cùng em "
    "trong suốt chặng đường vừa qua."
))

add_para(doc, (
    "Mặc dù đã dành nhiều tâm huyết và nỗ lực để hoàn thiện đồ án, nhưng do hạn chế về mặt thời gian và kinh nghiệm, "
    "sản phẩm chắc chắn không tránh khỏi những thiếu sót. Em rất mong nhận được những ý kiến đóng góp, phê bình "
    "quý báu từ quý Thầy Cô để em có thể rút kinh nghiệm và hoàn thiện kiến thức của mình hơn nữa trong tương lai."
))

add_para(doc, "Em xin chân thành cảm ơn!", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)


# ════════════════════════════════════════════════════════════════════
#  MỞ ĐẦU
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "MỞ ĐẦU", level=1)

# 1. Lý do chọn đề tài
add_heading(doc, "1. Lý do chọn đề tài", level=2)
add_para(doc, (
    "Trong kỷ nguyên chuyển đổi số, việc ứng dụng công nghệ vào quản lý vận hành du lịch đã trở thành "
    "yêu cầu cấp thiết. Các phương thức soát vé truyền thống hiện nay vẫn tồn tại nhiều bất cập như: "
    "tình trạng vé giả, việc chia sẻ mã QR không chính chủ, và khó khăn trong việc giám sát nhân viên "
    "tại các điểm kiểm soát. Công nghệ nhận diện khuôn mặt và ký số đang dần trở nên phổ biến, mang "
    "lại tính định danh duy nhất và độ bảo mật cao. Xuất phát từ nhu cầu thực tiễn đó, em quyết định "
    "thực hiện đề tài \"Hệ thống kiểm soát ra/vào khu du lịch đa kênh dựa trên QR và xác thực định danh\" "
    "nhằm tạo ra một giải pháp quản lý hiện đại, minh bạch và an toàn."
))

# 2. Mục tiêu nghiên cứu
add_heading(doc, "2. Mục tiêu nghiên cứu", level=2)
add_para(doc, (
    "Nghiên cứu và ứng dụng kỹ thuật nhận diện khuôn mặt SFace (State-of-the-art) để định danh khách hàng 1:1. "
    "Tìm hiểu cơ chế ký số RS256 để bảo mật vé điện tử. Xây dựng một hệ thống phân tán hoàn chỉnh gồm "
    "Mobile App cho nhân viên/khách hàng, Web Dashboard cho quản lý và AI Service cho xử lý sinh trắc học, "
    "đảm bảo quy trình soát vé diễn ra nhanh chóng dưới 3 giây."
))

# 3. Đối tượng và phạm vi nghiên cứu
add_heading(doc, "3. Đối tượng và phạm vi nghiên cứu", level=2)
add_para(doc, (
    "Đề tài tập trung nghiên cứu quy trình kiểm soát khách tham quan tại các cổng khu du lịch. "
    "Phạm vi thực nghiệm bao gồm việc triển khai hệ thống phần mềm đa nền tảng (Android, Web, Cloud) "
    "với 4 kênh xác thực chính: QR Code ký số, Nhận diện khuôn mặt, Tra cứu ID và Kiểm soát thủ công."
))

# 4. Ý nghĩa thực tiễn
add_heading(doc, "4. Ý nghĩa thực tiễn", level=2)
add_para(doc, (
    "Sản phẩm giúp tự động hóa khâu kiểm soát, giảm thiểu thất thoát doanh thu do gian lận vé. "
    "Đồng thời, hệ thống cung cấp dữ liệu phân tích vận hành trực quan cho ban quản lý thông qua "
    "Dashboard, hỗ trợ theo dõi hiệu suất làm việc của nhân viên và hành vi của khách hàng, "
    "từ đó nâng cao chất lượng dịch vụ và hình ảnh chuyên nghiệp của khu du lịch."
))

# 5. Tóm tắt nội dung
add_heading(doc, "5. Tóm tắt nội dung", level=2)
add_para(doc, "Nội dung đề tài được tổ chức thành 5 chương:")

# Chương 1
add_para(doc, "Chương 1: GIỚI THIỆU TỔNG QUAN", bold=True, indent=0.5)
add_para(doc, (
    "Trong chương này, em giới thiệu về bối cảnh thực tế, các vấn đề tồn tại của hệ thống soát vé cũ "
    "và đề xuất giải pháp xác thực đa kênh tích hợp AI."
), indent=1.0)

# Chương 2
add_para(doc, "Chương 2: CƠ SỞ LÝ THUYẾT", bold=True, indent=0.5)
add_para(doc, (
    "Tìm hiểu sâu về các công nghệ cốt lõi: Framework FastAPI, kiến trúc MongoDB NoSQL, thuật toán "
    "nhận diện khuôn mặt SFace và cơ chế ký số bảo mật RS256."
), indent=1.0)

# Chương 3
add_para(doc, "Chương 3: QUY TRÌNH HỆ THỐNG", bold=True, indent=0.5)
add_para(doc, (
    "Trình bày kiến trúc phân tán Cloud-Edge, thiết kế cơ sở dữ liệu và quy trình nghiệp vụ "
    "chi tiết từ khâu mua vé, đăng ký khuôn mặt đến soát vé tại cổng."
), indent=1.0)

# Chương 4
add_para(doc, "Chương 4: KẾT QUẢ VÀ THỰC NGHIỆM", bold=True, indent=0.5)
add_para(doc, (
    "Trình diễn giao diện ứng dụng Web/Mobile, kết quả kiểm thử tốc độ nhận diện AI và "
    "khả năng giám sát vận hành thông qua Dashboard."
), indent=1.0)

# Chương 5
add_para(doc, "Chương 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", bold=True, indent=0.5)
add_para(doc, (
    "Tổng kết các kết quả đạt được, những hạn chế còn tồn tại và đề xuất các hướng nâng cấp "
    "như công nghệ chống giả mạo (Anti-spoofing) và Edge AI."
), indent=1.0)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  II. KIẾN TRÚC HỆ THỐNG & TRIỂN KHAI
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "II. KIẾN TRÚC HỆ THỐNG & SƠ ĐỒ TRIỂN KHAI", level=1)

add_heading(doc, "2.1. Mô hình kiến trúc phân tán (Distributed Architecture)", level=2)
add_para(doc, (
    "Hệ thống được thiết kế theo mô hình phân tán gồm 4 tầng chính, tách biệt rõ ràng về "
    "trách nhiệm và cho phép mở rộng độc lập từng thành phần:"
))

layers = [
    ("Tầng Giao diện (Client)", "Android Gate App (nhân viên), Android Customer App (khách hàng), Web Dashboard (Admin/ReactJS)"),
    ("Tầng Máy chủ Đám mây (Cloud)", "FastAPI Backend chạy trên Render Cloud — xử lý nghiệp vụ, phân quyền RBAC, Audit Trail"),
    ("Tầng AI Chuyên biệt (Edge)", "AI Service (OpenCV SFace) chạy tại máy trạm cục bộ, kết nối Cloud qua đường hầm Ngrok"),
    ("Tầng Dữ liệu (Storage)", "MongoDB Atlas Cloud — lưu vé, danh tính, vector khuôn mặt, log hệ thống"),
]
add_table(doc,
    ["Tầng hệ thống", "Mô tả & Công nghệ"],
    layers
)

doc.add_paragraph()
add_heading(doc, "2.2. Cơ chế xác thực đa kênh (4 Authentication Channels)", level=2)
channels = [
    ("QR e-ticket", "Ký số RS256, chống giả mạo, cơ chế Nonce chống dùng lại vé", "Bắt buộc"),
    ("Face Verify 1:1", "So khớp khuôn mặt với embedding SFace (128-d) đã đăng ký", "Tùy chọn (opt-in)"),
    ("Identity Hash", "Tra cứu theo mã hash CCCD hoặc Booking ID", "Tùy chọn"),
    ("Manual Fallback", "Tra cứu thủ công theo SĐT / tên khách", "Dự phòng"),
]
add_table(doc,
    ["Kênh xác thực", "Cơ chế hoạt động", "Trạng thái"],
    channels
)

doc.add_paragraph()
add_heading(doc, "2.3. Sơ đồ triển khai thực tế", level=2)
add_para(doc, (
    "Backend được deploy tự động lên Render Cloud thông qua GitHub CI/CD. Web Dashboard "
    "chạy trên Cloudflare Pages (CDN toàn cầu). AI Service chạy cục bộ trên máy trạm có "
    "GPU/CPU mạnh và được expose ra Internet thông qua Ngrok Tunnel để Backend Cloud gọi được."
))
deploy_info = [
    ("Backend API", "Render Cloud", "https://check-ticket-1hyd.onrender.com"),
    ("Web Dashboard", "Cloudflare Pages", "https://fc439656.tourism-dashboard.pages.dev"),
    ("Database", "MongoDB Atlas (M0)", "Cluster khang1402 — region Asia"),
    ("AI Service", "Local + Ngrok Tunnel", "Port 8001 → expose qua Ngrok free tier"),
]
add_table(doc,
    ["Thành phần", "Nền tảng triển khai", "URL / Địa chỉ"],
    deploy_info
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  III. CÔNG NGHỆ & TÍNH NĂNG CHÍNH
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "III. CÔNG NGHỆ VÀ TÍNH NĂNG CHÍNH", level=1)

add_heading(doc, "3.1. Công nghệ sử dụng", level=2)
techs = [
    ("Android App", "Kotlin Native, Retrofit2, CameraX, ZXing (QR), Google ML Kit, Bottom Navigation"),
    ("Backend API", "FastAPI (Python), Motor async, JWT HS256, Pydantic v2, Uvicorn"),
    ("Web Dashboard", "ReactJS + Vite, Recharts (biểu đồ), Axios, Zustand (state)"),
    ("Cơ sở dữ liệu", "MongoDB Atlas Cloud (NoSQL, Free Tier M0)"),
    ("AI – Khuôn mặt", "OpenCV FaceRecognizerSF (SFace), vector 128-d, Cosine Similarity"),
    ("Bảo mật QR", "RS256 Digital Signature (python-jose), Nonce anti-reuse"),
    ("Realtime", "WebSocket — Staff Presence & Gate Event push"),
    ("Triển khai", "Render, Cloudflare Pages, Ngrok, GitHub Actions"),
]
add_table(doc,
    ["Thành phần", "Công nghệ cụ thể"],
    techs
)

doc.add_paragraph()
add_heading(doc, "3.2. Tính năng chính của hệ thống", level=2)

features = [
    ("🔐 Nhận diện khuôn mặt SFace",
     "Sử dụng model OpenCV SFace tối ưu hóa 128-d vector. Ngưỡng tương đồng tinh chỉnh "
     "ở mức 0.37 đảm bảo độ chính xác cao cho gương mặt người Việt Nam. Không lưu ảnh "
     "gốc — chỉ lưu vector (Privacy by Design)."),
    ("📱 Ứng dụng Mobile Dual-role",
     "App dành cho Nhân viên: quét QR, xác thực khuôn mặt tại cổng, bán vé tại quầy, "
     "tra cứu thủ công. App dành cho Khách hàng: đăng ký tài khoản, mua vé online, "
     "đăng ký Face ID, xem lịch sử vé và đánh giá 1–5 sao."),
    ("🌐 Web Dashboard vận hành",
     "Giao diện quản trị đa trang: Tổng quan realtime, Giám sát từng cổng, Thống kê "
     "doanh thu/lượt khách, Phân tích cơ cấu nhóm tuổi, Quản lý vé/khách hàng/nhân viên, "
     "Xem đánh giá, Xuất báo cáo CSV."),
    ("🔍 Audit Trail & Giám sát nhân viên",
     "Mọi request đều được ghi audit_log (user, IP, endpoint, status, thời gian). "
     "Tính năng Staff Presence theo dõi trạng thái Online/Offline của nhân viên thời gian "
     "thực qua WebSocket."),
    ("🤖 Trợ lý AI Chatbot",
     "Tích hợp Chatbot thông minh trên cả Web Dashboard và App Mobile. Hỗ trợ nhân viên "
     "và khách hàng giải đáp thắc mắc, tra cứu thông tin vé tự động 24/7."),
    ("🛡️ Bảo mật đa lớp",
     "JWT HS256 cho API, RS256 cho QR Code, RBAC phân quyền 4 vai trò (Admin/Manager/"
     "Operator/Cashier), Hash SHA-256 cho dữ liệu nhạy cảm (CCCD)."),
]

for title, desc in features:
    add_heading(doc, title, level=3, color=(0,70,127))
    add_para(doc, desc, indent=0.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  IV. KẾT QUẢ ĐẠT ĐƯỢC
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "IV. KẾT QUẢ ĐẠT ĐƯỢC", level=1)

add_heading(doc, "4.1. Kết quả kỹ thuật", level=2)
results = [
    ("Tốc độ xác thực", "~1.5 – 2.0 giây/lượt (bao gồm cả giao tiếp mạng Cloud ↔ AI)"),
    ("Độ chính xác AI", "Ngưỡng Cosine 0.37 — nhận diện ổn định trong điều kiện ánh sáng bình thường"),
    ("Hiệu suất Server", "Backend FastAPI async xử lý đồng thời nhiều request không nghẽn cổ chai"),
    ("Độ ổn định", "Hệ thống vận hành ổn định trên môi trường Cloud Render, DB Atlas 24/7"),
    ("Bảo mật QR", "RS256 + Nonce: chưa ghi nhận trường hợp vé giả qua được cổng kiểm soát"),
]
add_table(doc, ["Chỉ số", "Kết quả đo được"], results)

doc.add_paragraph()
add_heading(doc, "4.2. Hệ sinh thái đã hoàn thiện", level=2)
ecosystem = [
    "Backend API: 20+ endpoints REST bao phủ toàn bộ nghiệp vụ (Auth, Tickets, Face, Checkin, Reports, WebSocket).",
    "Web Dashboard: 10+ trang quản trị với biểu đồ Recharts trực quan, xuất CSV báo cáo.",
    "Android App: 15+ màn hình phục vụ 2 luồng người dùng (Nhân viên và Khách hàng).",
    "AI Service: Pipeline hoàn chỉnh Detect → Align → Extract (SFace 128-d) → Cosine Compare.",
    "Triển khai thực tế: Backend trên Render, Web trên Cloudflare Pages, DB trên MongoDB Atlas.",
]
for e in ecosystem:
    add_bullet(doc, e)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  V. PHÂN TÍCH & KẾT LUẬN
# ════════════════════════════════════════════════════════════════════
add_heading(doc, "V. PHÂN TÍCH HỆ THỐNG VÀ KẾT LUẬN", level=1)

add_heading(doc, "5.1. Ưu điểm nổi bật", level=2)
pros = [
    "Kiến trúc Microservices tinh gọn — AI Service tách rời Backend, dễ nâng cấp model AI mà không ảnh hưởng hệ thống.",
    "Đa kênh linh hoạt — Khách hàng có nhiều phương án dự phòng khi một kênh gặp sự cố.",
    "Privacy by Design — Tuân thủ nguyên tắc bảo vệ dữ liệu cá nhân ngay từ thiết kế.",
    "Audit Trail toàn diện — Hỗ trợ kiểm tra trách nhiệm và phát hiện bất thường.",
    "Realtime WebSocket — Giám sát vận hành cổng và nhân viên tức thời.",
]
for p_text in pros:
    add_bullet(doc, p_text)

add_heading(doc, "5.2. Hạn chế còn tồn tại", level=2)
cons = [
    "Tốc độ phản hồi AI phụ thuộc vào chất lượng đường truyền Internet (Cloud ↔ Local AI qua Ngrok).",
    "Chất lượng nhận diện giảm trong điều kiện ánh sáng yếu hoặc Camera điện thoại độ phân giải thấp.",
    "Ngrok Free Tier giới hạn băng thông và đổi URL mỗi lần khởi động lại — cần cập nhật thủ công.",
]
for c_text in cons:
    add_bullet(doc, c_text)

add_heading(doc, "5.3. Hướng phát triển", level=2)
futures = [
    "Tích hợp Liveness Detection (phát hiện chớp mắt/mỉm cười) để chống nạn dùng ảnh/video giả mạo khuôn mặt.",
    "Triển khai Edge AI — chạy model SFace trực tiếp trên chip thiết bị (Offline-ready), không phụ thuộc mạng.",
    "Tích hợp thanh toán online (VNPay, Momo) vào luồng mua vé của khách hàng.",
    "Mở rộng phân tích cảm xúc khách hàng (Emotion Analytics) để đánh giá mức độ hài lòng.",
    "Triển khai Kubernetes để tự động mở rộng quy mô (Auto-scaling) khi lượng khách tăng đột biến.",
]
for f_text in futures:
    add_bullet(doc, f_text)

add_heading(doc, "5.4. Kết luận", level=2)
add_para(doc, (
    "Đề tài đã xây dựng thành công một hệ thống kiểm soát ra/vào khu du lịch hoàn chỉnh, "
    "hiện đại và có tính thực tiễn cao. Với việc tích hợp AI nhận diện khuôn mặt SFace, "
    "cơ chế QR ký số RS256, kiến trúc Cloud phân tán và Dashboard vận hành đa chiều, hệ "
    "thống đáp ứng được bài toán kiểm soát an ninh thông minh cho các khu du lịch quy mô "
    "vừa và lớn. Sản phẩm đã được triển khai và kiểm thử thực tế trên môi trường Cloud, "
    "chứng minh tính khả thi và độ ổn định trong điều kiện vận hành thực tế."
))

# ════════════════════════════════════════════════════════════════════
#  TÀI LIỆU THAM KHẢO
# ════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1)
refs = [
    "[1] OpenCV SFace Model — https://github.com/opencv/opencv_zoo",
    "[2] FastAPI Documentation — https://fastapi.tiangolo.com",
    "[3] MongoDB Atlas — https://www.mongodb.com/atlas",
    "[4] Android CameraX — https://developer.android.com/camerax",
    "[5] Google ML Kit Face Detection — https://developers.google.com/ml-kit",
    "[6] Ngrok Tunneling — https://ngrok.com/docs",
    "[7] Render Cloud Platform — https://render.com/docs",
    "[8] InsightFace SFace Paper — ICCV 2021",
]
for ref in refs:
    add_bullet(doc, ref)

# ─── Lưu file ────────────────────────────────────────────────────────
out_path = r"e:\Learn\Do_an_nganh\Check_ticket\DO_AN_TOT_NGHIEP_NGO_MANH_KHANG.docx"
doc.save(out_path)
print("DONE: " + out_path)
